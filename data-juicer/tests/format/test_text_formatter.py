import io
import os
import tempfile
import unittest
from types import SimpleNamespace
from unittest.mock import patch

from cryptography.fernet import Fernet

from data_juicer.format.load import load_formatter
from data_juicer.format.text_formatter import TextFormatter, _decrypt_and_extract
from data_juicer.utils.unittest_utils import DataJuicerTestCaseBase


def _make_global_cfg(key_path, decrypt=True):
    """Helper: return a minimal cfg namespace understood by load_dataset."""
    return SimpleNamespace(
        decrypt_after_reading=decrypt,
        encryption_key_path=key_path,
    )


class TextFormatterTest(DataJuicerTestCaseBase):

    def setUp(self):
        super().setUp()

        self._path = os.path.join(os.path.dirname(os.path.realpath(__file__)), "data", "text")
        self._file = os.path.join(self._path, "sample1.txt")

    def test_text_file(self):
        formatter = TextFormatter(self._file)
        ds = formatter.load_dataset()
        self.assertEqual(len(ds), 1)
        self.assertIn("text", list(ds.features.keys()))

    def test_text_path(self):
        formatter = TextFormatter(self._path)
        ds = formatter.load_dataset()
        self.assertEqual(len(ds), 6)
        self.assertIn("text", list(ds.features.keys()))

    def test_text_path_with_suffixes(self):
        formatter = TextFormatter(self._path, suffixes=[".txt"])
        ds = formatter.load_dataset()
        self.assertEqual(len(ds), 6)
        self.assertIn("text", list(ds.features.keys()))

    def test_text_path_with_add_suffix(self):
        formatter = TextFormatter(self._path, add_suffix=True)
        ds = formatter.load_dataset()
        self.assertEqual(len(ds), 6)
        self.assertIn("text", list(ds.features.keys()))
        self.assertIn("__dj__suffix__", list(ds.features.keys()))

    def test_load_formatter_with_file(self):
        """Test load_formatter with a direct text file path"""
        formatter = load_formatter(self._file)
        self.assertIsInstance(formatter, TextFormatter)
        ds = formatter.load_dataset()
        self.assertEqual(len(ds), 1)
        self.assertIn("text", list(ds.features.keys()))

    def test_load_formatter_with_specified_suffix(self):
        """Test load_formatter with specified suffixes"""
        formatter = load_formatter(self._path, suffixes=[".txt"])
        self.assertIsInstance(formatter, TextFormatter)
        ds = formatter.load_dataset()
        self.assertEqual(len(ds), 6)
        self.assertIn("text", list(ds.features.keys()))


# ---------------------------------------------------------------------------
# Tests for the decrypt_after_reading multiprocessing path
# ---------------------------------------------------------------------------

class TextFormatterDecryptTest(DataJuicerTestCaseBase):
    """Tests for the encrypted-read path added to TextFormatter.

    These tests exercise:
    - ``_decrypt_and_extract``: the top-level worker function used by Pool
    - ``TextFormatter.load_dataset`` with a ``global_cfg`` that enables
      ``decrypt_after_reading`` for plain .txt files
    - ``TextFormatter.load_dataset`` with a ``global_cfg`` that enables
      ``decrypt_after_reading`` for .docx files (parallel decrypt+extract)
    """

    def setUp(self):
        super().setUp()
        # Generate a fresh Fernet key for every test so tests are independent.
        self.fernet = Fernet.generate_key()
        self.f = Fernet(self.fernet)
        self._text_path = os.path.join(
            os.path.dirname(os.path.realpath(__file__)), "data", "text"
        )

    # ------------------------------------------------------------------
    # Helper utilities
    # ------------------------------------------------------------------

    def _write_key_file(self, tmp_dir):
        key_path = os.path.join(tmp_dir, "test.key")
        with open(key_path, "wb") as fh:
            fh.write(self.fernet)
        return key_path

    def _encrypt_text_file(self, src_path, dst_dir):
        """Return path to encrypted copy of *src_path* in *dst_dir*."""
        fname = os.path.basename(src_path)
        dst_path = os.path.join(dst_dir, fname)
        with open(src_path, "rb") as fh:
            plaintext = fh.read()
        with open(dst_path, "wb") as fh:
            fh.write(self.f.encrypt(plaintext))
        return dst_path

    # ------------------------------------------------------------------
    # _decrypt_and_extract (worker function)
    # ------------------------------------------------------------------

    def test_decrypt_and_extract_docx(self):
        """_decrypt_and_extract correctly decrypts a .docx and writes .txt."""
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx not installed")

        with tempfile.TemporaryDirectory() as tmp:
            # Create a minimal .docx in memory
            doc = Document()
            doc.add_paragraph("Hello from encrypted docx")
            src_docx = os.path.join(tmp, "sample.docx")
            doc.save(src_docx)

            # Encrypt the docx
            enc_docx = os.path.join(tmp, "sample_enc.docx")
            with open(src_docx, "rb") as fh:
                plaintext = fh.read()
            with open(enc_docx, "wb") as fh:
                fh.write(self.f.encrypt(plaintext))

            out_dir = os.path.join(tmp, "out")
            os.makedirs(out_dir)

            # Run the worker — must not raise and must produce a .txt file
            _decrypt_and_extract(enc_docx, self.fernet, out_dir, ".docx")

            txt_files = [f for f in os.listdir(out_dir) if f.endswith(".txt")]
            self.assertEqual(len(txt_files), 1)
            with open(os.path.join(out_dir, txt_files[0])) as fh:
                content = fh.read()
            self.assertIn("Hello from encrypted docx", content)

    def test_decrypt_and_extract_pdf(self):
        """_decrypt_and_extract correctly decrypts a .pdf and writes .txt."""
        try:
            import pdfplumber
            from reportlab.pdfgen import canvas as rl_canvas
        except ImportError:
            self.skipTest("pdfplumber or reportlab not installed")

        with tempfile.TemporaryDirectory() as tmp:
            src_pdf = os.path.join(tmp, "sample.pdf")
            c = rl_canvas.Canvas(src_pdf)
            c.drawString(100, 750, "Hello from encrypted pdf")
            c.save()

            enc_pdf = os.path.join(tmp, "sample_enc.pdf")
            with open(src_pdf, "rb") as fh:
                plaintext = fh.read()
            with open(enc_pdf, "wb") as fh:
                fh.write(self.f.encrypt(plaintext))

            out_dir = os.path.join(tmp, "out")
            os.makedirs(out_dir)

            _decrypt_and_extract(enc_pdf, self.fernet, out_dir, ".pdf")

            txt_files = [f for f in os.listdir(out_dir) if f.endswith(".txt")]
            self.assertEqual(len(txt_files), 1)

    def test_decrypt_and_extract_wrong_key_raises(self):
        """_decrypt_and_extract raises InvalidToken for wrong key."""
        from cryptography.fernet import InvalidToken

        with tempfile.TemporaryDirectory() as tmp:
            try:
                from docx import Document
            except ImportError:
                self.skipTest("python-docx not installed")

            doc = Document()
            doc.add_paragraph("secret")
            src_docx = os.path.join(tmp, "sample.docx")
            doc.save(src_docx)

            enc_docx = os.path.join(tmp, "sample_enc.docx")
            with open(src_docx, "rb") as fh:
                plaintext = fh.read()
            with open(enc_docx, "wb") as fh:
                fh.write(self.f.encrypt(plaintext))

            wrong_key = Fernet.generate_key()
            out_dir = os.path.join(tmp, "out")
            os.makedirs(out_dir)

            with self.assertRaises(InvalidToken):
                _decrypt_and_extract(enc_docx, wrong_key, out_dir, ".docx")

    # ------------------------------------------------------------------
    # TextFormatter.load_dataset — plain .txt decrypt path
    # ------------------------------------------------------------------

    def test_load_dataset_decrypt_txt_files(self):
        """load_dataset decrypts .txt files in-memory and loads them correctly."""
        with tempfile.TemporaryDirectory() as tmp:
            key_path = self._write_key_file(tmp)

            # Encrypt all sample .txt files into tmp/
            enc_dir = os.path.join(tmp, "enc_txt")
            os.makedirs(enc_dir)
            src_files = [
                os.path.join(self._text_path, f)
                for f in os.listdir(self._text_path)
                if f.endswith(".txt")
            ]
            for src in src_files:
                self._encrypt_text_file(src, enc_dir)

            formatter = TextFormatter(enc_dir, suffixes=[".txt"])
            global_cfg = _make_global_cfg(key_path)
            ds = formatter.load_dataset(num_proc=1, global_cfg=global_cfg)

            self.assertEqual(len(ds), len(src_files))
            self.assertIn("text", list(ds.features.keys()))

    def test_load_dataset_decrypt_txt_content_matches_plaintext(self):
        """Decrypted content loaded by load_dataset matches original plaintext."""
        with tempfile.TemporaryDirectory() as tmp:
            key_path = self._write_key_file(tmp)

            # Use a single known-content file
            sample_path = os.path.join(self._text_path, "sample4.txt")
            enc_dir = os.path.join(tmp, "enc")
            os.makedirs(enc_dir)
            self._encrypt_text_file(sample_path, enc_dir)

            formatter = TextFormatter(enc_dir, suffixes=[".txt"])
            global_cfg = _make_global_cfg(key_path)
            ds = formatter.load_dataset(num_proc=1, global_cfg=global_cfg)

            # Load the plaintext version for comparison
            plain_formatter = TextFormatter(sample_path)
            plain_ds = plain_formatter.load_dataset(num_proc=1)

            self.assertEqual(ds[0]["text"], plain_ds[0]["text"])

    def test_load_dataset_decrypt_false_skips_encryption(self):
        """When decrypt_after_reading=False, plaintext files load normally."""
        formatter = TextFormatter(self._text_path, suffixes=[".txt"])
        global_cfg = _make_global_cfg(key_path=None, decrypt=False)
        ds = formatter.load_dataset(num_proc=1, global_cfg=global_cfg)
        self.assertEqual(len(ds), 6)

    # ------------------------------------------------------------------
    # TextFormatter.load_dataset — .docx parallel decrypt path
    # ------------------------------------------------------------------

    def test_load_dataset_decrypt_docx_multiproc(self):
        """load_dataset runs _decrypt_and_extract in parallel for .docx files."""
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx not installed")

        with tempfile.TemporaryDirectory() as tmp:
            key_path = self._write_key_file(tmp)

            enc_dir = os.path.join(tmp, "enc_docx")
            os.makedirs(enc_dir)

            # Create and encrypt two minimal .docx files
            expected_texts = []
            for i in range(2):
                doc = Document()
                msg = f"Paragraph {i} content"
                doc.add_paragraph(msg)
                expected_texts.append(msg)
                plain_path = os.path.join(tmp, f"doc{i}.docx")
                doc.save(plain_path)
                enc_path = os.path.join(enc_dir, f"doc{i}.docx")
                with open(plain_path, "rb") as fh:
                    plaintext = fh.read()
                with open(enc_path, "wb") as fh:
                    fh.write(self.f.encrypt(plaintext))

            formatter = TextFormatter(enc_dir, suffixes=[".docx"])
            global_cfg = _make_global_cfg(key_path)

            # Verify Pool is actually called (parallel execution)
            with patch("data_juicer.format.text_formatter.Pool") as mock_pool_cls:
                mock_pool = mock_pool_cls.return_value.__enter__ = mock_pool_cls.return_value
                mock_pool.apply_async = unittest.mock.MagicMock()
                mock_pool.close = unittest.mock.MagicMock()
                mock_pool.join = unittest.mock.MagicMock()
                # We only verify Pool is instantiated; actual extraction tested separately
                try:
                    formatter.load_dataset(num_proc=2, global_cfg=global_cfg)
                except Exception:
                    pass  # May fail due to mock; we care Pool was created
                mock_pool_cls.assert_called_once_with(2)


if __name__ == "__main__":
    unittest.main()
